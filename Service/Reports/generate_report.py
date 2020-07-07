from docx import Document
from docx.shared import Inches
from datetime import date


def generate():
    document = Document()

    document.add_heading('Blame Report', 0)

    p = document.add_paragraph("Ok ok, I'll change the name. This is an example paragraph"
                               " with any information you want to add statically to the report."
                               " I can format elements in  ")
    p.add_run('bold').bold = True
    p.add_run(' and ')
    p.add_run('italic.').italic = True

    document.add_heading('Report', level=1)
    document.add_paragraph('Problem type being reported', style='Intense Quote')

    document.add_paragraph(
        'First issue being reported', style='List Bullet'
    )
    document.add_paragraph(
        'Second issue being reported', style='List Bullet'
    )
    document.add_paragraph(
        'Proof pictures can be added here', style='List Number'
    )

    document.add_picture('3.png', width=Inches(3.75))

    records = (
        ('John Doof', '22345678', 'Mill uncleaned'),
        ('Jane Dirty', '21354126', 'unreported tool breakage'),
        ('Innocent Accused', '20967347', 'Unmatched signin')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Student Id'
    hdr_cells[2].text = 'Reason for report'
    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    document.save(f"Service/Reports/{date.today()}_report.docx")



